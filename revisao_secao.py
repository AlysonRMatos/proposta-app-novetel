"""Monta a secao "Revisao" do documento (subdocumento docxtpl), usada
apenas quando a proposta atual e uma revisao de uma proposta existente.
Quando nao e revisao, o subdoc fica vazio e a secao nao aparece."""
from docx.shared import Emu, Pt

from itens_tabela import montar_tabela_itens

FONT_NAME = "Trebuchet MS"
HEADING_COLOR_HEX = "1F497D"

# Mesmo recuo padrao usado no restante do documento (ex: secao "Quem
# somos?"), para os paragrafos desta secao nao ficarem colados na margem.
RECUO_ESQUERDO_PADRAO = Emu(540385)
RECUO_PRIMEIRA_LINHA_PADRAO = Emu(374015)


def montar_secao_revisao(
    subdoc,
    ativa: bool,
    codigo_pai: str = None,
    numero_revisao: int = None,
    solicitacao_alteracao: str = None,
    itens_antigos: list = None,
    itens_novos: list = None,
    itens_alterados: list = None,
    itens_adicionados: list = None,
    justificativas_itens: dict = None,
):
    if not ativa:
        return

    justificativas_itens = justificativas_itens or {}

    from docx.shared import RGBColor

    def add_heading(texto, size=20):
        p = subdoc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(texto)
        run.font.name = "Arial Narrow"
        run.font.size = Pt(size)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(HEADING_COLOR_HEX)

    def add_body(texto, bold=False, space_after=8):
        p = subdoc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = RECUO_ESQUERDO_PADRAO
        p.paragraph_format.first_line_indent = RECUO_PRIMEIRA_LINHA_PADRAO
        run = p.add_run(texto)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.bold = bold

    add_heading(f"Revisão RV{numero_revisao:02d}")
    if codigo_pai:
        add_body(f"Referente à proposta original: {codigo_pai}", space_after=10)

    def add_bullet(texto):
        p = subdoc.add_paragraph()
        p.paragraph_format.left_indent = RECUO_ESQUERDO_PADRAO
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"•  {texto}")
        run.font.name = FONT_NAME
        run.font.size = Pt(11)

    add_body("Solicitação de alteração desta revisão:", bold=True, space_after=4)
    add_body(solicitacao_alteracao or "Não informado.", space_after=14)

    if itens_alterados:
        add_body("Itens com quantidade alterada nesta revisão:", bold=True, space_after=4)
        for item in itens_alterados:
            motivo = justificativas_itens.get(item["codigo"]) or "Motivo não informado."
            add_bullet(
                f"{item['codigo']} - {item['descricao']}: "
                f"{item['qtd_antiga']} → {item['qtd_nova']} {item['unidade']} — Motivo: {motivo}"
            )
        subdoc.add_paragraph()

    if itens_adicionados:
        add_body("Itens novos incluídos nesta revisão:", bold=True, space_after=4)
        for item in itens_adicionados:
            motivo = justificativas_itens.get(item["codigo"]) or "Motivo não informado."
            add_bullet(
                f"{item['codigo']} - {item['descricao']} "
                f"({item['quantidade']} {item['unidade']}) — Motivo: {motivo}"
            )
        subdoc.add_paragraph()

    if itens_antigos:
        add_body("Itens da versão anterior (referência):", bold=True, space_after=6)
        montar_tabela_itens(subdoc, itens_antigos)
        subdoc.add_paragraph()

    if itens_novos:
        add_body("Itens desta revisão (atual):", bold=True, space_after=6)
        montar_tabela_itens(subdoc, itens_novos)
