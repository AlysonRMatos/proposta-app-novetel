"""
Gera templates/template_proposta.docx a partir do documento ORIGINAL enviado
pelo usuario, preservando 100% da formatacao (fontes, cores, logos, marca
d'agua, cabecalho/rodape). Apenas os trechos variaveis sao trocados por
placeholders Jinja (docxtpl).

Rodar uma unica vez (ou sempre que o modelo original mudar):
    python build_template.py
"""
import copy
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

# Recuo padrao usado nos paragrafos de texto corrido do documento original
# (ex: secao "Quem somos?"), para os textos que adicionamos seguirem o
# mesmo padrao visual em vez de ficarem colados na margem.
RECUO_ESQUERDO_PADRAO = Emu(540385)
RECUO_PRIMEIRA_LINHA_PADRAO = Emu(374015)


def aplicar_recuo_padrao(paragraph):
    paragraph.paragraph_format.left_indent = RECUO_ESQUERDO_PADRAO
    paragraph.paragraph_format.first_line_indent = RECUO_PRIMEIRA_LINHA_PADRAO

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(BASE_DIR, "templates", "template_proposta.docx")

ORIGINAL_SOURCE = (
    r"C:\Users\alyso\OneDrive\Área de Trabalho\NOVETEL\TRABALHO\Propostas\SHOPEE\Propostas"
    r"\Proposta_Técnica_Comercial_HUB-LRJ-11 - ELETRICA - JIRA - INFRA - 1623-2504.RV01"
    r"\Proposta_Técnica_Comercial_HUB-LRJ-11 - ELETRICA - JIRA - INFRA - 1623-2504.RV01.docx"
)


def clear_and_set_text(paragraph, new_text):
    """Remove todos os runs (inclusive imagens/drawings) e escreve um unico
    run novo de texto simples, preservando a formatacao (fonte, tamanho, cor)
    do primeiro run original."""
    runs = list(paragraph.runs)
    rPr_original = None
    for run in runs:
        rPr = run._element.find(qn("w:rPr"))
        if rPr_original is None and rPr is not None:
            rPr_original = copy.deepcopy(rPr)
        run._element.getparent().remove(run._element)

    if new_text:
        new_run = paragraph.add_run(new_text)
        if rPr_original is not None:
            existing_rPr = new_run._element.find(qn("w:rPr"))
            if existing_rPr is not None:
                new_run._element.remove(existing_rPr)
            new_run._element.insert(0, rPr_original)


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def mover_bookmark_start(paragraph_origem, paragraph_destino, bookmark_id):
    """Move o <w:bookmarkStart> de id=bookmark_id do inicio de paragraph_origem
    para o inicio de paragraph_destino. Necessario quando o paragrafo de
    origem sera removido/substituido, mas o <w:bookmarkEnd> correspondente
    (em outro paragrafo, referenciado por um PAGEREF no indice) precisa
    continuar tendo um bookmarkStart valido em algum lugar antes dele."""
    p_origem = paragraph_origem._p
    p_destino = paragraph_destino._p
    for el in p_origem.findall(qn("w:bookmarkStart")):
        if el.get(qn("w:id")) == str(bookmark_id):
            p_origem.remove(el)
            p_destino.insert(0, el)
            return True
    return False


def main():
    if not os.path.exists(ORIGINAL_SOURCE):
        raise FileNotFoundError(f"Documento original nao encontrado: {ORIGINAL_SOURCE}")

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc = Document(ORIGINAL_SOURCE)
    paragraphs = doc.paragraphs

    # ---------- Capa ----------
    clear_and_set_text(paragraphs[4], "Instalações {{ disciplina }}")
    clear_and_set_text(
        paragraphs[7],
        "{{ cliente }} | {{ codigo_projeto }} - {{ local }} | {{ escopo_titulo }}",
    )
    clear_and_set_text(paragraphs[12], "{{ data_proposta }}.")

    # ---------- Objetivos ----------
    clear_and_set_text(paragraphs[70], "{{ cliente }}.")
    clear_and_set_text(paragraphs[72], "Endereço: {{ endereco }}.")
    clear_and_set_text(paragraphs[73], "Cidade: {{ cidade }}.")
    clear_and_set_text(paragraphs[76], "Objeto: {{ objeto }}")
    clear_and_set_text(
        paragraphs[83],
        "Na oportunidade, informamos que temos total interesse na prestação "
        "dos serviços junto à {{ cliente }}.",
    )

    # ---------- Documentos disponibilizados ----------
    # paragrafo 99 contem a foto especifica da proposta original -> vira o
    # ponto de insercao dinamica dos anexos de cada nova proposta.
    clear_and_set_text(paragraphs[99], "{{p imagens}}")
    # paragrafos 100 a 135 sao espacadores vazios que reservavam altura de
    # pagina para as fotos fixas do modelo original; somados, empurravam o
    # conteudo e geravam uma pagina em branco antes de "Escopo". Removidos
    # pois o grid dinamico de imagens ja controla seu proprio espaco.
    for p in paragraphs[100:136]:
        remove_paragraph(p)

    # "Documentos disponibilizados" deve ficar isolado (somente as fotos);
    # "Escopo" sempre comeca em pagina nova, tenha poucas ou muitas fotos.
    paragraphs[136].paragraph_format.page_break_before = True

    # ---------- Especificações ----------
    # o paragrafo 170 contem o inicio do bookmark _Toc190957976, cujo fim
    # esta no paragrafo 178 ("Prazo e Preço") -- usado pelo indice (PAGEREF).
    # Precisa ser preservado antes de apagar o paragrafo, senao o indice
    # gera "Erro! Indicador nao definido." para essa entrada.
    mover_bookmark_start(paragraphs[170], paragraphs[167], bookmark_id=6)

    # paragrafos 170/172/174/176 descreviam os itens da proposta ORIGINAL em
    # texto livre; substituimos por uma tabela dinamica vinda da LPU.
    clear_and_set_text(paragraphs[170], "{{p itens_tabela}}")

    # Secao de observacoes (itens exclusos da proposta), logo apos a tabela.
    obs_titulo = paragraphs[172].insert_paragraph_before()
    obs_titulo.add_run("Observações - itens exclusos desta proposta:").bold = True
    aplicar_recuo_padrao(obs_titulo)
    obs_texto = paragraphs[172].insert_paragraph_before()
    obs_texto.add_run("{{ observacoes_exclusao }}")
    aplicar_recuo_padrao(obs_texto)

    # Secao de Revisao (so aparece quando a proposta e uma revisao de uma
    # proposta existente; subdoc vazio nao mostra nada caso contrario).
    revisao_p = paragraphs[172].insert_paragraph_before()
    revisao_p.add_run("{{p secao_revisao}}")

    remove_paragraph(paragraphs[172])
    remove_paragraph(paragraphs[174])
    remove_paragraph(paragraphs[176])

    # ---------- Prazo e Preço ----------
    clear_and_set_text(
        paragraphs[182],
        "O total previsto para execução do projeto é de {{ prazo_execucao }} "
        "a partir do aceite da proposta e mobilização da equipe e materiais.",
    )
    clear_and_set_text(
        paragraphs[189],
        "O valor total da presente proposta, considerando todos os encargos, "
        "tributos e obrigações para a correta execução do serviço, conforme "
        "a boa norma construtiva perfazem o montante de {{ valor_total_extenso }}.",
    )

    # ---------- Atualizacao automatica de campos (indice/PAGEREF) ----------
    # Sem isso, os numeros de pagina do indice ficam "congelados" com os
    # valores em cache do documento original, ja que nem python-docx nem
    # LibreOffice recalculam layout ao salvar via codigo. Esta flag manda o
    # Word (e o LibreOffice, que a respeita) recalcular todos os campos
    # automaticamente ao abrir/converter o documento.
    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.append(update_fields)

    # ---------- Numero da proposta (canto superior da capa) ----------
    primeiro_paragrafo = doc.paragraphs[0]
    p_codigo = primeiro_paragrafo.insert_paragraph_before()
    p_codigo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_codigo.paragraph_format.right_indent = Cm(1.2)
    run_codigo = p_codigo.add_run("Nº {{ codigo_proposta }}")
    run_codigo.font.size = Pt(9)
    run_codigo.font.name = "Trebuchet MS"

    doc.save(OUT_PATH)
    print(f"Template gerado a partir do documento original em: {OUT_PATH}")


if __name__ == "__main__":
    main()
