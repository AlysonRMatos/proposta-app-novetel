"""Monta um grid padronizado de imagens (2 colunas, tamanho uniforme, com
margem entre elas) para a secao "Documentos disponibilizados"."""
from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLUNAS = 2
LARGURA_COLUNA = Inches(3.05)
LARGURA_IMAGEM = Inches(2.75)
ESPACO_CELULA = Pt(6)


def _remover_bordas_tabela(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def _definir_margem_celulas(table, margem=ESPACO_CELULA):
    tbl_pr = table._tbl.tblPr
    cell_mar = OxmlElement("w:tblCellMar")
    for edge in ("top", "start", "bottom", "end"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(margem.twips))
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tbl_pr.append(cell_mar)


def montar_grid_imagens(subdoc, caminhos_imagens, legendas=None):
    if not caminhos_imagens:
        p = subdoc.add_paragraph("(Nenhum documento anexado)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    legendas = legendas or [None] * len(caminhos_imagens)
    n = len(caminhos_imagens)
    linhas = (n + COLUNAS - 1) // COLUNAS

    table = subdoc.add_table(rows=linhas, cols=COLUNAS)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _remover_bordas_tabela(table)
    _definir_margem_celulas(table)

    for idx, caminho in enumerate(caminhos_imagens):
        r, c = divmod(idx, COLUNAS)
        cell = table.rows[r].cells[c]
        cell.width = LARGURA_COLUNA

        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(caminho, width=LARGURA_IMAGEM)

        legenda = legendas[idx]
        if legenda:
            p_legenda = cell.add_paragraph(legenda)
            p_legenda.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_legenda in p_legenda.runs:
                run_legenda.font.size = Pt(8)
                run_legenda.font.name = "Trebuchet MS"

    # preenche celulas vazias da ultima linha (caso numero impar de imagens)
    resto = (COLUNAS - (n % COLUNAS)) % COLUNAS
    for i in range(resto):
        cell = table.rows[linhas - 1].cells[COLUNAS - 1 - i]
        cell.width = LARGURA_COLUNA
