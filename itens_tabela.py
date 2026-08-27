"""Monta a tabela de especificacoes (itens da LPU) como subdocumento docxtpl."""
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT_NAME = "Trebuchet MS"
HEADER_BG = "1F497D"
FONT_SIZE = Pt(9)

# Larguras redistribuidas: Codigo/Especificacao precisam de mais espaco,
# Qnt./Unidade sao sempre curtos (numeros e siglas), entao ficam estreitos.
LARGURAS = [Inches(1.2), Inches(3.9), Inches(0.6), Inches(0.7)]
CABECALHOS = ["Codigo", "Especificacao", "Qnt.", "Unidade"]


def _set_cell_bg(cell, color_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def montar_tabela_itens(subdoc, itens):
    table = subdoc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Sem isso, o Word ignora as larguras definidas abaixo e ajusta cada
    # coluna pelo conteudo (foi o que deixava "Qnt." larga demais).
    table.autofit = False

    header_cells = table.rows[0].cells
    for i, h in enumerate(CABECALHOS):
        header_cells[i].width = LARGURAS[i]
        _set_cell_bg(header_cells[i], HEADER_BG)
        p = header_cells[i].paragraphs[0]
        if i >= 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    if not itens:
        row = table.add_row()
        cell = row.cells[0]
        for c in row.cells[1:]:
            c.merge(cell)
        cell.text = "Nenhum item selecionado."
        return table

    for item in itens:
        row = table.add_row()
        valores = [
            item.get("codigo", ""),
            item.get("descricao", ""),
            str(item.get("quantidade", "")),
            item.get("unidade", ""),
        ]
        for i, val in enumerate(valores):
            cell = row.cells[i]
            cell.width = LARGURAS[i]
            p = cell.paragraphs[0]
            if i >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
    return table
